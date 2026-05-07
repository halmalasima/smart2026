from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


DOCX = Path(r"D:\المستوى الرابع\ترم ثاني\مشروع تخرج 2 الوهباني\مشروع التخرج - منقح ومفهرس.docx")
WORK_DOCX = Path(r"D:\smart2026\graduation_project_clean_indexed.docx")
OVERALL_PNG = Path(r"D:\smart2026\use_case_overall_smartjudi.png")
ACCOUNTS_PNG = Path(r"D:\smart2026\use_case_accounts_smartjudi.png")

FONT = r"C:\Windows\Fonts\arial.ttf"
BOLD_FONT = r"C:\Windows\Fonts\arialbd.ttf"


def draw_rtl(draw, xy, text, font, fill=(20, 20, 20), anchor="mm"):
    try:
        draw.text(xy, text, font=font, fill=fill, anchor=anchor, direction="rtl")
    except TypeError:
        draw.text(xy, text, font=font, fill=fill, anchor=anchor)


def text_bbox(draw, text, font):
    try:
        return draw.textbbox((0, 0), text, font=font, direction="rtl")
    except TypeError:
        return draw.textbbox((0, 0), text, font=font)


def draw_centered_wrapped(draw, center, text, font, max_width, fill=(20, 20, 20)):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        bbox = text_bbox(draw, candidate, font)
        if bbox[2] - bbox[0] <= max_width or not line:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    line_h = font.size + 8
    y = center[1] - (len(lines) - 1) * line_h / 2
    for item in lines:
        draw_rtl(draw, (center[0], y), item, font, fill=fill)
        y += line_h


def draw_actor(draw, x, y, label, font, color=(25, 72, 55)):
    draw.ellipse((x - 18, y - 68, x + 18, y - 32), outline=color, width=4)
    draw.line((x, y - 32, x, y + 28), fill=color, width=4)
    draw.line((x - 40, y - 2, x + 40, y - 2), fill=color, width=4)
    draw.line((x, y + 28, x - 30, y + 72), fill=color, width=4)
    draw.line((x, y + 28, x + 30, y + 72), fill=color, width=4)
    draw_rtl(draw, (x, y + 112), label, font)


def draw_ellipse_case(draw, center, text, font, w=300, h=90, color=(37, 105, 78)):
    x, y = center
    box = (x - w / 2, y - h / 2, x + w / 2, y + h / 2)
    draw.ellipse(box, fill="white", outline=color, width=4)
    draw_centered_wrapped(draw, center, text, font, max_width=w - 60)


def draw_service_box(draw, xyxy, title, subtitle, title_font, sub_font):
    x1, y1, x2, y2 = xyxy
    draw.rounded_rectangle(xyxy, radius=22, outline=(22, 82, 58), width=4, fill=(250, 252, 250))
    draw.rounded_rectangle((x1 + 8, y1 + 8, x2 - 8, y1 + 54), radius=16, outline=(22, 82, 58), width=2, fill=(232, 242, 236))
    draw_rtl(draw, ((x1 + x2) / 2, y1 + 31), title, title_font, fill=(22, 82, 58))
    draw_centered_wrapped(draw, ((x1 + x2) / 2, (y1 + y2) / 2 + 25), subtitle, sub_font, max_width=(x2 - x1 - 50), fill=(45, 45, 45))


def line(draw, p1, p2, color=(115, 115, 115), width=3):
    draw.line((p1[0], p1[1], p2[0], p2[1]), fill=color, width=width)


def make_overall_diagram():
    img = Image.new("RGB", (1800, 1160), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(BOLD_FONT, 40)
    head_font = ImageFont.truetype(BOLD_FONT, 24)
    body_font = ImageFont.truetype(FONT, 22)
    actor_font = ImageFont.truetype(BOLD_FONT, 26)
    note_font = ImageFont.truetype(FONT, 20)

    draw_rtl(draw, (900, 50), "المخطط العام لحالات الاستخدام - منصة القضاء الذكي", title_font, fill=(22, 82, 58))
    draw.rounded_rectangle((250, 100, 1550, 1085), radius=30, outline=(22, 82, 58), width=5, fill=(246, 250, 247))
    draw_rtl(draw, (900, 132), "حدود النظام: SmartJudi Platform", note_font, fill=(22, 82, 58))

    actors = {
        "المسؤول": (110, 230),
        "المحامي": (110, 450),
        "المواطن": (110, 680),
        "القاضي": (1690, 230),
        "كاتب العدل": (1690, 450),
        "المعاون": (1690, 680),
    }
    for name, pos in actors.items():
        draw_actor(draw, pos[0], pos[1], name, actor_font)

    services = {
        "خدمة الحسابات والصلاحيات": ((380, 190, 770, 395), "تسجيل الدخول\nإدارة الحساب\nالحسابات الفرعية\nالصلاحيات"),
        "خدمة القضايا والدعاوى": ((835, 190, 1225, 395), "رفع الدعوى\nإدارة الأطراف\nتفاصيل القضية\nالطعون وأوامر الأداء"),
        "خدمة الجلسات والأحكام": ((1290, 190, 1495, 395), "جدولة الجلسات\nمتابعة الجلسات\nإصدار الأحكام"),
        "خدمة البحث والمكتبة القانونية": ((380, 470, 770, 675), "القوانين والمواد\nالبحث القانوني\nالمحاكم والمحامون"),
        "خدمة المساعد الذكي": ((835, 470, 1225, 675), "الاستفسارات القانونية\nتحليل القضية\nRAG"),
        "خدمة الإشعارات والمراسلات": ((1290, 470, 1495, 675), "الإشعارات\nالرسائل\nسجل المحادثات"),
        "خدمة المستندات والأرشفة": ((520, 750, 910, 955), "رفع المرفقات\nملف القضية\nالأرشفة"),
        "خدمة العقود والوكالات والمحاسبة": ((980, 750, 1370, 955), "العقود والوكالات\nالمحاسبة\nالخدمات التوثيقية"),
    }
    for title, (box, subtitle) in services.items():
        draw_service_box(draw, box, title, subtitle, head_font, body_font)

    assoc = [
        ("المسؤول", "خدمة الحسابات والصلاحيات"),
        ("المسؤول", "خدمة الإشعارات والمراسلات"),
        ("المحامي", "خدمة الحسابات والصلاحيات"),
        ("المحامي", "خدمة القضايا والدعاوى"),
        ("المحامي", "خدمة البحث والمكتبة القانونية"),
        ("المحامي", "خدمة المساعد الذكي"),
        ("المحامي", "خدمة المستندات والأرشفة"),
        ("المواطن", "خدمة الحسابات والصلاحيات"),
        ("المواطن", "خدمة القضايا والدعاوى"),
        ("المواطن", "خدمة البحث والمكتبة القانونية"),
        ("المواطن", "خدمة المساعد الذكي"),
        ("القاضي", "خدمة الجلسات والأحكام"),
        ("القاضي", "خدمة القضايا والدعاوى"),
        ("القاضي", "خدمة البحث والمكتبة القانونية"),
        ("كاتب العدل", "خدمة العقود والوكالات والمحاسبة"),
        ("كاتب العدل", "خدمة الحسابات والصلاحيات"),
        ("المعاون", "خدمة المستندات والأرشفة"),
        ("المعاون", "خدمة القضايا والدعاوى"),
    ]
    for actor, service in assoc:
        ax, ay = actors[actor]
        x1, y1, x2, y2 = services[service][0]
        target = ((x1 + x2) / 2, (y1 + y2) / 2)
        if ax < 900:
            start = (ax + 52, ay)
            end = (x1, target[1])
        else:
            start = (ax - 52, ay)
            end = (x2, target[1])
        line(draw, start, end, width=2)

    draw.rounded_rectangle((400, 1010, 1400, 1065), radius=16, fill=(232, 242, 236), outline=(120, 170, 140), width=2)
    draw_rtl(draw, (900, 1038), "يعرض هذا الشكل خدمات النظام الرئيسية فقط. وسيتم تفصيل كل خدمة في مخطط فرعي مستقل.", note_font, fill=(40, 85, 65))
    img.save(OVERALL_PNG)


def make_accounts_diagram():
    img = Image.new("RGB", (1700, 1040), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(BOLD_FONT, 38)
    font = ImageFont.truetype(FONT, 24)
    actor_font = ImageFont.truetype(BOLD_FONT, 26)
    note_font = ImageFont.truetype(FONT, 20)

    draw_rtl(draw, (850, 48), "مخطط فرعي: خدمة الحسابات والصلاحيات", title_font, fill=(22, 82, 58))
    draw.rounded_rectangle((290, 100, 1410, 970), radius=28, outline=(22, 82, 58), width=5, fill=(247, 251, 248))
    draw_rtl(draw, (850, 130), "حدود الخدمة: Account & Access Control Service", note_font, fill=(22, 82, 58))

    actors = {
        "مستخدم النظام": (110, 235),
        "المحامي": (110, 470),
        "المسؤول": (1590, 235),
        "المعاون": (1590, 470),
    }
    for name, pos in actors.items():
        draw_actor(draw, pos[0], pos[1], name, actor_font)

    cases = {
        "تسجيل الدخول": (560, 220),
        "التحقق عبر OTP": (850, 220),
        "إدارة الملف الشخصي": (1140, 220),
        "تغيير كلمة المرور": (560, 430),
        "تسجيل الخروج": (850, 430),
        "إنشاء حساب فرعي": (1140, 430),
        "إدارة المستخدمين": (690, 665),
        "إدارة الصلاحيات والأدوار": (1010, 665),
        "مراجعة سجلات الدخول": (850, 845),
    }
    for text, pos in cases.items():
        draw_ellipse_case(draw, pos, text, font, w=270, h=84)

    assoc = [
        ("مستخدم النظام", "تسجيل الدخول"),
        ("مستخدم النظام", "التحقق عبر OTP"),
        ("مستخدم النظام", "إدارة الملف الشخصي"),
        ("مستخدم النظام", "تغيير كلمة المرور"),
        ("مستخدم النظام", "تسجيل الخروج"),
        ("المحامي", "إنشاء حساب فرعي"),
        ("المعاون", "تسجيل الدخول"),
        ("المسؤول", "إدارة المستخدمين"),
        ("المسؤول", "إدارة الصلاحيات والأدوار"),
        ("المسؤول", "مراجعة سجلات الدخول"),
    ]
    for actor, case in assoc:
        ax, ay = actors[actor]
        ux, uy = cases[case]
        if ax < 850:
            start = (ax + 52, ay)
            end = (ux - 135, uy)
        else:
            start = (ax - 52, ay)
            end = (ux + 135, uy)
        line(draw, start, end, width=2)

    # include relations
    include_color = (70, 120, 100)
    line(draw, (560 + 115, 220), (850 - 115, 220), color=include_color, width=2)
    draw_rtl(draw, (850, 192), "<<include>>", note_font, fill=include_color)
    line(draw, (1140 - 130, 430 + 8), (1010 + 130, 665 - 8), color=include_color, width=2)
    draw_rtl(draw, (1165, 560), "يشرف عليها المسؤول", note_font, fill=include_color, anchor="rm")

    draw.rounded_rectangle((470, 910, 1230, 955), radius=14, fill=(232, 242, 236), outline=(120, 170, 140), width=2)
    draw_rtl(draw, (850, 932), "الحسابات الفرعية يقصد بها ربط المساعد أو المستخدم التابع بالمحامي أو الجهة الإدارية.", note_font, fill=(40, 85, 65))
    img.save(ACCOUNTS_PNG)


def set_rtl(paragraph, center=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_font(run, size=14, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)


def remove_generated_sections(doc):
    triggers = {
        "3.7.1 المخطط العام لحالات الاستخدام",
        "3.7.2 مخطط خدمة الحسابات والصلاحيات",
        "3.7.1 مخطط حالات الاستخدام Use Case Diagram",
    }
    removing = False
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if text in triggers:
            removing = True
        elif removing and p.style.name.startswith("Heading") and text.startswith("3.7.") and text not in triggers:
            removing = False
        if removing:
            el = p._element
            el.getparent().remove(el)


def add_heading_after(doc, paragraph, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    set_rtl(p)
    insert_after(paragraph, p)
    return p


def add_text_after(doc, paragraph, text, center=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=13)
    set_rtl(p, center=center)
    insert_after(paragraph, p)
    return p


def add_image_after(doc, paragraph, image_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(17.2))
    insert_after(paragraph, p)
    return p


def add_caption_after(doc, paragraph, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=12, bold=True)
    set_rtl(p, center=True)
    insert_after(paragraph, p)
    return p


def update_docx():
    doc = Document(WORK_DOCX if WORK_DOCX.exists() else DOCX)
    remove_generated_sections(doc)

    target = None
    for p in doc.paragraphs:
        if "3.7 المخططات التحليلية" in p.text:
            target = p
            break
    if target is None:
        raise RuntimeError("Section 3.7 not found")

    intro = add_text_after(
        doc,
        target,
        "يبدأ هذا القسم بمخطط عام يوضح الخدمات الرئيسية في النظام وعلاقتها بالمستخدمين، ثم تتبعه مخططات فرعية تفصيلية لكل خدمة حتى تكون القراءة أوضح والربط مع المشروع الفعلي أسهل.",
    )

    h1 = add_heading_after(doc, intro, "3.7.1 المخطط العام لحالات الاستخدام")
    d1 = add_text_after(
        doc,
        h1,
        "يوضح هذا المخطط الصورة العامة للنظام على مستوى الخدمات الرئيسية، مثل الحسابات، القضايا، الجلسات، البحث القانوني، المساعد الذكي، الإشعارات، المستندات، والخدمات التوثيقية.",
    )
    i1 = add_image_after(doc, d1, OVERALL_PNG)
    c1 = add_caption_after(doc, i1, "الشكل (3-1): المخطط العام لحالات الاستخدام لمنصة القضاء الذكي.")

    h2 = add_heading_after(doc, c1, "3.7.2 مخطط خدمة الحسابات والصلاحيات")
    d2 = add_text_after(
        doc,
        h2,
        "يوضح هذا المخطط تفاصيل خدمة الحسابات والصلاحيات، بما يشمل تسجيل الدخول، التحقق عبر OTP، إدارة الملف الشخصي، الحسابات الفرعية، إدارة المستخدمين، وإدارة الأدوار والصلاحيات.",
    )
    i2 = add_image_after(doc, d2, ACCOUNTS_PNG)
    add_caption_after(doc, i2, "الشكل (3-2): مخطط حالات الاستخدام لخدمة الحسابات والصلاحيات.")

    doc.save(WORK_DOCX)


if __name__ == "__main__":
    make_overall_diagram()
    make_accounts_diagram()
    update_docx()
    print(DOCX)
    print(OVERALL_PNG)
    print(ACCOUNTS_PNG)
