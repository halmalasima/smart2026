from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"D:\smart2026\graduation_project_clean_indexed.docx"


def set_run_font(run, size=14, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = rPr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            rPr.append(rtl)
        rtl.set(qn("w:val"), "1")


def add_para(doc, text="", style=None, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    set_rtl(p, align)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size={1: 18, 2: 16, 3: 14}.get(level, 14), bold=True, color=(26, 86, 64))
    set_rtl(p)
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, size=13)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if p.runs:
        p.runs[0].text = "• " + p.runs[0].text
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    set_rtl(p)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=12)
        shade_cell(table.rows[0].cells[i], "E7F0EA")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=11)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if widths and i < len(widths):
                cell.width = Cm(widths[i])
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(14)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Arial"
        styles[name]._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    # Cover
    add_para(doc, "وزارة التربية والتعليم العالي والبحث العلمي", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "الجامعة اليمنية", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "كلية الهندسة وعلوم الحاسوب", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "منصة القضاء الذكي", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "SmartJudi Platform", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "إعداد الطلاب", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "محمد صالح محمد المنبه", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "أحمد محمود الدوبحي", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "محمد عبدالإله الشامي", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "إشراف الدكتور", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "عبدالحميد الوهباني", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(
        doc,
        "تم إنجاز هذا المشروع كجزء من متطلبات التخرج لنيل درجة البكالوريوس من كلية الهندسة وعلوم الحاسوب - قسم تقنية المعلومات.",
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "للعام الجامعي 2025م - 2026م", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    add_heading(doc, "ملخص المشروع", 1)
    add_para(doc, "يهدف هذا المشروع إلى تطوير منصة قضائية إلكترونية ذكية تسهم في دعم التحول الرقمي للمنظومة القضائية في اليمن، وذلك من خلال تقديم خدمات قانونية متكاملة تسهّل إجراءات التقاضي وتحسّن كفاءة الوصول إلى المعلومات القانونية.")
    add_para(doc, "تسعى المنصة إلى معالجة مشكلات بطء الإجراءات الورقية، وصعوبة متابعة القضايا والجلسات، وتشتت المصادر القانونية، من خلال بيئة رقمية موحدة تجمع بين إدارة القضايا، والمكتبة القانونية، ودليل المحامين، والإشعارات، والمساعد القانوني الذكي.")
    add_para(doc, "يعتمد المشروع في تنفيذه على Flutter وDart لبناء واجهات الهواتف والويب، وعلى Django وDjango REST Framework لبناء النظام الخلفي، مع استخدام FastAPI في خدمات الذكاء الاصطناعي ومحرك RAG. كما يعتمد على PostgreSQL للبيانات الأساسية، وChromaDB كمستودع متجهات للبحث الذكي، مع دعم SQLite في بيئة التطوير المحلية.")
    add_para(doc, "ومن المتوقع أن يسهم النظام في تسريع إجراءات التقاضي، ورفع كفاءة البحث القانوني، وتحسين الشفافية، وتقليل الاعتماد على المعاملات الورقية، مع التأكيد أن المساعد الذكي يقدم دعماً معلوماتياً ولا يصدر أحكاماً قانونية نهائية.")

    add_heading(doc, "Project Abstract", 1)
    add_para(doc, "This project aims to develop an intelligent electronic judicial platform that supports the digital transformation of the Yemeni judicial system by providing integrated legal services for case management, legal research, notifications, and AI-assisted legal support.", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "The implemented system uses Flutter and Dart for the mobile and web frontend, Django and Django REST Framework for the backend APIs, FastAPI for AI-related services, PostgreSQL for core data storage, and ChromaDB for vector-based legal retrieval. The AI assistant is designed to support legal search and analysis without replacing official judicial decision-making.", align=WD_ALIGN_PARAGRAPH.LEFT)
    page_break(doc)

    add_heading(doc, "التفويض", 1)
    add_para(doc, "نحن الطلبة الموقعين أدناه، نقر بأن هذا المشروع تم إنجازه من قبلنا وبجهد ذاتي، وبإشراف ومتابعة المشرف الأكاديمي، وذلك لاستكمال متطلبات الحصول على درجة البكالوريوس في تخصص تقنية المعلومات من كلية الهندسة وعلوم الحاسوب.")
    add_table(doc, ["م", "اسم الطالب", "التوقيع", "التاريخ"], [
        ["1", "محمد صالح محمد المنبه", "", ""],
        ["2", "أحمد محمود الدوبحي", "", ""],
        ["3", "محمد عبدالإله الشامي", "", ""],
    ], widths=[1.2, 7, 3, 3])

    add_heading(doc, "الإهداء", 1)
    add_para(doc, "إلى آبائنا وأمهاتنا الذين كانوا سنداً لنا في مسيرتنا العلمية، وإلى إخوتنا وأصدقائنا وكل من وقف بجانبنا ودعمنا حتى وصلنا إلى هذه المرحلة، نهدي هذا العمل المتواضع عرفاناً وتقديراً.")
    add_para(doc, "وإلى دكاترتنا الكرام الذين قدموا لنا العلم والتوجيه، وكانوا مصدر إلهام لنا في بناء هذا المشروع.")

    add_heading(doc, "الشكر والتقدير", 1)
    add_para(doc, "نحمد الله عز وجل الذي وفقنا في إتمام هذا المشروع. كما نتقدم بجزيل الشكر والتقدير إلى مشرف المشروع الدكتور عبدالحميد الوهباني على توجيهاته وملاحظاته، وإلى جميع أعضاء هيئة التدريس في كلية الهندسة وعلوم الحاسوب على ما قدموه لنا من علم ودعم خلال سنوات الدراسة.")

    add_heading(doc, "شهادة المشرف", 1)
    add_para(doc, "أشهد بأن مشروع التخرج المعنون بـ «منصة القضاء الذكي» قد أعده الطلاب: محمد صالح محمد المنبه، أحمد محمود الدوبحي، محمد عبدالإله الشامي، تحت إشرافي في قسم تقنية المعلومات، وذلك كجزء من متطلبات الحصول على درجة البكالوريوس.")
    add_para(doc, "اسم المشرف: د/ عبدالحميد الوهباني")
    add_para(doc, "التوقيع: ____________________")
    add_para(doc, "التاريخ: ____________________")

    add_heading(doc, "لجنة المناقشة", 1)
    add_table(doc, ["الصفة", "الاسم", "التوقيع"], [
        ["المشرف", "د/ عبدالحميد الوهباني", ""],
        ["عضو لجنة المناقشة", "", ""],
        ["عضو لجنة المناقشة", "", ""],
        ["رئيس القسم", "د/ نصر الدين الماوري", ""],
    ], widths=[4, 6, 4])
    page_break(doc)

    add_heading(doc, "فهرس المحتويات", 1)
    toc_rows = [
        ["ملخص المشروع"],
        ["Project Abstract"],
        ["التفويض"],
        ["الإهداء"],
        ["الشكر والتقدير"],
        ["الفصل الأول: الدراسة التمهيدية"],
        ["1.1 نظرة عامة"],
        ["1.2 بيان المشكلة"],
        ["1.3 أهداف المشروع"],
        ["1.4 أهمية المشروع"],
        ["1.5 نطاق المشروع"],
        ["1.6 منهجية العمل"],
        ["1.7 النتائج المتوقعة"],
        ["الفصل الثاني: الإطار النظري والدراسات السابقة"],
        ["2.1 الخلفية النظرية"],
        ["2.2 الدراسات السابقة"],
        ["2.3 فوائد القضاء الإلكتروني"],
        ["الفصل الثالث: تحليل وتوصيف النظام"],
        ["3.1 المقدمة"],
        ["3.2 الجهات المستفيدة"],
        ["3.3 المتطلبات الوظيفية وغير الوظيفية"],
        ["3.4 حالات الاستخدام النصية"],
        ["3.5 مكونات النظام"],
        ["3.6 سيناريو العمل المتكامل"],
        ["3.7 المخططات التحليلية - المرحلة الثانية"],
    ]
    add_table(doc, ["العنوان"], toc_rows, widths=[14])
    page_break(doc)

    add_heading(doc, "الفصل الأول: الدراسة التمهيدية", 1)
    add_heading(doc, "1.1 نظرة عامة", 2)
    add_para(doc, "يشهد النظام القضائي في اليمن حاجة متزايدة إلى حلول رقمية تساعد على تسهيل الإجراءات، وتنظيم البيانات، وتحسين الوصول إلى الخدمات القانونية. ومن هنا جاءت فكرة منصة القضاء الذكي بوصفها منصة رقمية تجمع بين تطبيقات الهواتف الذكية والويب، ونظام خلفي لإدارة البيانات، وخدمات ذكاء اصطناعي مساعدة للبحث القانوني وتحليل المعلومات.")
    add_para(doc, "تسعى المنصة إلى تقديم تجربة موحدة للمواطن والمحامي والقاضي والموثق ومدير النظام، بحيث يتمكن كل مستخدم من الوصول إلى الخدمات المناسبة لصلاحياته عبر واجهات عربية واضحة وسهلة الاستخدام.")

    add_heading(doc, "1.2 بيان المشكلة", 2)
    add_para(doc, "يعاني العمل القضائي التقليدي من بطء الإجراءات، وكثرة الاعتماد على المعاملات الورقية، وصعوبة متابعة القضايا والجلسات، إضافة إلى تشتت مصادر القوانين والنماذج القانونية. كما يواجه المواطنون والمحامون صعوبة في الوصول السريع إلى المعلومات القانونية المناسبة، مما يؤدي إلى زيادة الوقت والتكلفة وتعقيد متابعة الإجراءات.")

    add_heading(doc, "1.3 أهداف المشروع", 2)
    for item in [
        "تصميم وتطوير منصة إلكترونية ذكية تخدم أطراف المنظومة القضائية عبر تطبيق Flutter للهواتف والويب.",
        "إدارة القضايا والدعاوى والجلسات والأحكام والطعون وأوامر الأداء من خلال نظام رقمي موحد.",
        "توفير مكتبة قانونية رقمية للقوانين اليمنية مع إمكانية البحث المتقدم.",
        "توفير مساعد ذكي يعتمد على تقنيات RAG وNLP لدعم الاستفسارات القانونية وتحليل القضايا.",
        "توفير دليل للمحامين والمحاكم وخدمات إشعارات ومراسلات داخلية مرتبطة بالنظام.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.4 أهمية المشروع", 2)
    for item in [
        "تسريع إجراءات التقاضي وتقليل الاعتماد على المعاملات الورقية.",
        "تسهيل وصول المواطنين والمحامين إلى المعلومات القانونية والخدمات القضائية.",
        "رفع كفاءة البحث القانوني من خلال المكتبة القانونية والمساعد الذكي.",
        "تعزيز الشفافية عبر سجلات العمليات والإشعارات ومتابعة حالة القضايا.",
        "تقديم أساس قابل للتطوير لإضافة خدمات قضائية جديدة مستقبلاً.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.5 نطاق المشروع", 2)
    add_para(doc, "يركز المشروع على القوانين والخدمات القضائية اليمنية، ويشمل إدارة المستخدمين، القضايا، الدعاوى، الأطراف، الجلسات، الأحكام، الطعون، المرفقات، المكتبة القانونية، دليل المحامين، الإشعارات، وخدمات الذكاء الاصطناعي القانونية.")
    add_para(doc, "لا يصدر المساعد الذكي أحكاماً قانونية نهائية، وإنما يقدم دعماً معلوماتياً وتحليلياً يساعد المستخدم في الوصول إلى النصوص القانونية وفهمها.")

    add_heading(doc, "1.6 منهجية العمل", 2)
    add_table(doc, ["المرحلة", "الوصف"], [
        ["تحليل المتطلبات", "دراسة احتياجات المستخدمين وتحديد وظائف النظام وصلاحيات الأدوار."],
        ["جمع البيانات", "تجميع القوانين والبيانات المرجعية مثل المحاكم والمحامين والنصوص القانونية."],
        ["التصميم", "تصميم بنية النظام، قواعد البيانات، الواجهات، وواجهات API."],
        ["التطوير", "بناء الواجهة باستخدام Flutter وبناء الخلفية باستخدام Django/FastAPI."],
        ["الاختبار", "فحص وظائف النظام الأساسية مثل الدخول، القضايا، البحث، والإشعارات."],
        ["النشر", "تجهيز النظام للتشغيل المحلي أو الاستضافة السحابية عبر Render وDocker."],
    ], widths=[4, 10])

    add_heading(doc, "1.7 النتائج المتوقعة", 2)
    for item in [
        "توفير منصة رقمية موحدة للخدمات القضائية والقانونية.",
        "تقليل زمن تنفيذ الإجراءات مقارنة بالطرق التقليدية.",
        "تحسين الوصول إلى القوانين والنماذج القانونية.",
        "تمكين المستخدمين من متابعة القضايا والجلسات والإشعارات بصورة أسهل.",
        "بناء نظام قابل للتوسع والتطوير المستقبلي.",
    ]:
        add_bullet(doc, item)
    page_break(doc)

    add_heading(doc, "الفصل الثاني: الإطار النظري والدراسات السابقة", 1)
    add_heading(doc, "2.1 الخلفية النظرية", 2)
    add_heading(doc, "2.1.1 القضاء الإلكتروني", 3)
    add_para(doc, "القضاء الإلكتروني هو استخدام تقنيات المعلومات والاتصالات في إدارة إجراءات التقاضي، وتبادل الوثائق، وتنظيم القضايا والجلسات، وتسهيل الوصول إلى الخدمات القانونية. ويسهم هذا الاتجاه في تقليل الزمن والجهد وتحسين دقة البيانات وسهولة الرجوع إليها.")
    add_heading(doc, "2.1.2 الذكاء الاصطناعي في المجال القانوني", 3)
    add_para(doc, "يساعد الذكاء الاصطناعي في المجال القانوني على تحسين البحث في النصوص القانونية، وتصنيف المعلومات، وتقديم إجابات مساعدة للمستخدمين. وفي هذا المشروع يتم استخدام فكرة الاسترجاع المعزز بالتوليد RAG لربط المساعد الذكي بمصادر قانونية داخلية بدلاً من الاعتماد على النموذج اللغوي وحده.")
    add_heading(doc, "2.1.3 تقنية RAG", 3)
    add_para(doc, "تعتمد تقنية RAG على استرجاع النصوص ذات العلاقة من قاعدة معرفة أو مستودع متجهات، ثم استخدامها في توليد إجابة أكثر ارتباطاً بالسياق. في منصة القضاء الذكي يتم توظيف هذه التقنية للبحث في النصوص القانونية اليمنية ومساعدة المستخدم في الوصول إلى المواد ذات الصلة.")

    add_heading(doc, "2.2 الدراسات السابقة", 2)
    add_para(doc, "تشير التجارب المرتبطة بالقضاء الإلكتروني إلى أن رقمنة الإجراءات القضائية تساعد في تحسين الكفاءة وتقليل التكاليف وتسهيل وصول المستفيدين إلى المعلومات. كما توضح الدراسات الحديثة في الذكاء الاصطناعي القانوني أن محركات البحث الذكية والمساعدات المعتمدة على اللغة الطبيعية يمكن أن تدعم المحامين والباحثين في الوصول السريع إلى النصوص القانونية.")
    add_para(doc, "وبالنظر إلى الواقع اليمني، فإن وجود مبادرات للتحول الرقمي في وزارة العدل والمحاكم يمثل أساساً يمكن البناء عليه لتطوير حلول أكثر تكاملاً مثل منصة القضاء الذكي.")

    add_heading(doc, "2.3 فوائد القضاء الإلكتروني", 2)
    for item in [
        "توفير الوقت والجهد في متابعة القضايا والإجراءات.",
        "تقليل مخاطر فقدان الوثائق الورقية من خلال الأرشفة الرقمية.",
        "تسهيل وصول الأطراف إلى المعلومات القانونية والقضائية.",
        "تعزيز الشفافية ومتابعة الإجراءات من خلال سجلات إلكترونية.",
        "تحسين تجربة المستخدم عبر واجهات رقمية واضحة.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.4 المراجع الأولية", 2)
    refs = [
        "وزارة العدل وحقوق الإنسان اليمنية، بوابة الخدمات والأخبار العدلية.",
        "تقارير ومبادرات التحول الرقمي في القضاء اليمني.",
        "مراجع تقنية حول نظم استرجاع المعلومات وRAG والذكاء الاصطناعي القانوني.",
        "مستندات المشروع الداخلية الخاصة بمحرك RAG وخدمات SmartJudi.",
    ]
    for r in refs:
        add_bullet(doc, r)
    page_break(doc)

    add_heading(doc, "الفصل الثالث: تحليل وتوصيف النظام", 1)
    add_heading(doc, "3.1 المقدمة", 2)
    add_para(doc, "يتناول هذا الفصل تحليل متطلبات منصة القضاء الذكي وتوصيف مكوناتها الأساسية. وقد تم الاعتماد في هذا التوصيف على فكرة المشروع وعلى المكونات البرمجية الموجودة فعلياً في النظام، مثل تطبيق Flutter، وخدمات Django، ومحرك الذكاء الاصطناعي، ونماذج البيانات الخاصة بالقضايا والمستخدمين والجلسات.")

    add_heading(doc, "3.2 الجهات المستفيدة", 2)
    add_table(doc, ["الدور", "المسؤوليات"], [
        ["المسؤول", "إدارة النظام والمستخدمين والصلاحيات وسجلات العمليات."],
        ["القاضي", "استعراض القضايا والجلسات والأحكام حسب الصلاحيات المتاحة."],
        ["المحامي", "رفع الدعاوى، إدارة القضايا، متابعة الجلسات، واستخدام البحث والمساعد الذكي."],
        ["المواطن", "متابعة الخدمات والقضايا المرتبطة به واستقبال الإشعارات."],
        ["كاتب العدل", "استخدام الخدمات التوثيقية والمالية المرتبطة بطبيعة الدور."],
        ["المعاون", "مساعدة المحامي في الأرشفة والمهام الإدارية تحت إشرافه."],
    ], widths=[3.5, 10.5])

    add_heading(doc, "3.3 متطلبات النظام", 2)
    add_heading(doc, "3.3.1 المتطلبات الوظيفية", 3)
    add_table(doc, ["المجموعة", "المتطلبات"], [
        ["إدارة الحسابات", "تسجيل الدخول، JWT، حفظ الجلسة، الأدوار والصلاحيات، الحسابات الفرعية."],
        ["إدارة القضايا", "إنشاء القضايا والدعاوى، إدارة الأطراف، الحالة، الطلبات، الأرشفة."],
        ["الجلسات والأحكام", "عرض الجلسات، جدولة المواعيد، إضافة الملاحظات، وربط الأحكام بالقضايا."],
        ["المكتبة القانونية", "تصفح القوانين والمواد والبحث في النصوص القانونية اليمنية."],
        ["الذكاء الاصطناعي", "محادثة قانونية، سجل محادثات، وربط بمحرك RAG وخدمات Groq/LLM."],
        ["الإشعارات والمراسلات", "إشعارات داخل التطبيق ومراسلات مرتبطة بالمستخدمين والقضايا."],
        ["الملفات", "رفع المرفقات والمستندات وربطها بملف القضية."],
    ], widths=[4, 10])

    add_heading(doc, "3.3.2 المتطلبات غير الوظيفية", 3)
    for item in [
        "الأمان: استخدام JWT وصلاحيات مبنية على الدور لمنع الوصول غير المصرح به.",
        "سهولة الاستخدام: واجهات عربية متوافقة مع الهواتف والويب.",
        "الأداء: تحسين زمن استجابة API والبحث القانوني.",
        "قابلية التوسع: إمكانية الانتقال من Monolith إلى Microservices عبر API Gateway.",
        "الموثوقية: تنظيم البيانات والمرفقات وسجلات العمليات لضمان الرجوع إليها.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.4 حالات الاستخدام النصية", 2)
    add_heading(doc, "3.4.1 رفع دعوى إلكترونية", 3)
    add_para(doc, "الفاعل الرئيسي: المحامي أو المواطن حسب الصلاحية.")
    for item in [
        "تسجيل الدخول إلى النظام.",
        "اختيار خدمة رفع دعوى أو إدارة القضايا.",
        "إدخال بيانات القضية والمحكمة والأطراف.",
        "رفع المستندات والمرفقات الداعمة.",
        "حفظ الدعوى ومتابعة حالتها من لوحة المستخدم.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.4.2 استشارة المساعد الذكي", 3)
    add_para(doc, "الفاعل الرئيسي: المحامي أو المواطن.")
    for item in [
        "الدخول إلى شاشة المساعد الذكي.",
        "كتابة السؤال القانوني بلغة طبيعية.",
        "استرجاع النصوص ذات الصلة من قاعدة المعرفة القانونية.",
        "عرض الإجابة للمستخدم مع مراعاة أن الإجابة مساعدة وليست حكماً قانونياً.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.5 مكونات النظام", 2)
    add_table(doc, ["المكون", "الوصف"], [
        ["واجهة المستخدم", "تطبيق Flutter للهواتف والويب، يدعم العربية ويدير الحالة باستخدام Provider."],
        ["واجهة API", "Django REST Framework لتقديم خدمات المصادقة والقضايا والقوانين والإشعارات."],
        ["الذكاء الاصطناعي", "FastAPI، Groq/LLM، RAG، LangChain، HuggingFace Embeddings، ChromaDB."],
        ["قاعدة البيانات", "PostgreSQL للإنتاج، SQLite للتطوير المحلي، وجداول للنظام القضائي والقانوني."],
        ["البنية التحتية", "Render، Docker، Nginx Gateway في مسار الخدمات المصغرة."],
    ], widths=[4, 10])

    add_heading(doc, "3.6 سيناريو العمل المتكامل", 2)
    add_para(doc, "يبدأ المستخدم بتسجيل الدخول إلى منصة القضاء الذكي، ثم تظهر له الواجهة المناسبة لصلاحياته. يستطيع المحامي إنشاء دعوى جديدة وإدخال بيانات الأطراف ورفع المستندات، ثم يتابع حالة القضية والجلسات من خلال النظام. ويمكن للقاضي أو المستخدم المخول استعراض تفاصيل القضية والجلسات والأحكام وفق الصلاحيات المحددة.")
    add_para(doc, "خلال سير العمل، يقوم النظام بإرسال إشعارات داخلية عند وجود تحديثات مهمة، مثل تحديد جلسة أو تحديث حالة القضية. كما يستطيع المستخدم استخدام المساعد الذكي للبحث في القوانين اليمنية وطرح أسئلة قانونية تساعده في فهم النصوص ذات العلاقة.")

    add_heading(doc, "3.7 المخططات التحليلية - المرحلة الثانية", 2)
    add_para(doc, "سيتم إعداد المخططات التحليلية في المرحلة الثانية من العمل بصورة مشتركة، حتى تكون مطابقة للمشروع الفعلي وقاعدة البيانات والواجهات المنفذة. وتشمل هذه المخططات:")
    for item in [
        "مخطط حالات الاستخدام Use Case Diagram.",
        "مخطط تدفق البيانات DFD Level 0 وDFD Level 1.",
        "مخطط الأصناف Class Diagram.",
        "مخطط التسلسل Sequence Diagram.",
        "مخطط الكيانات والعلاقات ERD.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "ملاحظة تنظيمية", 3)
    add_para(doc, "تم تنظيف التقرير حتى نقطة المخططات. في المرحلة التالية سيتم بناء المخططات بناءً على ملفات المشروع الفعلية مثل نماذج Django، مسارات API، وشاشات Flutter، ثم إدراجها في هذا الفصل.")

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
